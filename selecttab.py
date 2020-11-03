from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import sys
from docx import Document


class SelectTab(QWidget):
    def __init__(self, link):
        super().__init__()
        self.database = link.database
        self.tree_model = link.tree_model
        self.signal = link.send_signal
        self.tree_view = None
        self.signal.connect(self.update_tree)
        self.initializeTab()

    def initializeTab(self):
        select_tab = QVBoxLayout()

        print_groupbox = self.initialize_print_box()
        tree_groupbox = self.make_select_tree_view()
        # buttons_groupbox = self.initialize_select_button_box()

        select_tab.addWidget(print_groupbox)
        select_tab.addWidget(tree_groupbox)
        # select_tab.addWidget(buttons_groupbox)

        self.setLayout(select_tab)

    def initialize_print_box(self):
        hbox = QHBoxLayout()

        checkboxes = self.initialize_check_boxes()
        export_button = QPushButton("Eksportuj jako docx")
        export_button.clicked.connect(self.export_tree_as_docx)

        hbox.addWidget(checkboxes)
        hbox.addWidget(export_button)

        print_box = QGroupBox()
        print_box.setLayout(hbox)
        return print_box

    def initialize_check_boxes(self):
        vbox = QVBoxLayout()
        checkbox1 = QCheckBox("zaznacz wszyskto")
        checkbox1.toggled.connect(self.checkAll)
        checkbox2 = QCheckBox("checkbox 2")
        # checkbox2.setChecked(True)

        vbox.addWidget(checkbox1)
        vbox.addWidget(checkbox2)

        groupbox = QGroupBox()
        groupbox.setLayout(vbox)
        return groupbox

    def export_tree_as_docx(self):
        root = self.tree_model.invisibleRootItem()
        files = SelectTab.gather_files_from_tree(root)
        Print_docx(files)

    @staticmethod
    def gather_files_from_tree(root):
        def children_of_(item):
            for index in range(item.rowCount()):
                name_col = item.child(index, 0)
                text_col = item.child(index, 1)
                yield name_col, text_col

        lst = []
        for name_item, text_item in children_of_(root):
            is_partially_checked = name_item.checkState() != Qt.Unchecked
            if is_partially_checked:
                if name_item.hasChildren():
                    files_to_print = SelectTab.gather_files_from_tree(name_item)
                    lst.extend(files_to_print)
                else:
                    file_to_print = text_item.text()
                    lst.append(file_to_print)
        return lst

    def make_select_tree_view(self):
        tree_view = QTreeView()
        tree_view.setModel(self.tree_model)
        tree_view.header().setSectionResizeMode(QHeaderView.ResizeToContents)
        tree_view.setEditTriggers(QAbstractItemView.NoEditTriggers)
        tree_view.hideColumn(1)
        tree_view.setHeaderHidden(True)
        tree_view.clicked.connect(self.on_item_clicked)
        self.tree_view = tree_view
        return tree_view

    def update_tree(self):
        self.tree_view.hideColumn(1)
        self.tree_view.expandAll()
        self.tree_view.setHeaderHidden(True)

    def checkAll(self, checkbox_state):
        root = self.tree_model.invisibleRootItem()
        if checkbox_state:
            self.check_all_descendants(root)
        else:
            self.uncheck_all_descendants(root)

    def on_item_clicked(self, index):
        item = self.tree_model.itemFromIndex(index)
        if item.checkState() == Qt.Checked:
            self.check_all_descendants(item)
            self.update_all_ancestors_Checked(item)
        elif item.checkState() == Qt.Unchecked:
            self.uncheck_all_descendants(item)
            self.update_all_ancestors_Unhecked(item)

    def update_all_ancestors_Checked(self, item):
        parent = item.parent()
        if parent is None:
            return

        tristate = 1
        for child in SelectTab.children_of_(parent):
            if not child.checkState() == Qt.Checked:
                parent.setCheckState(tristate)
                self.update_all_ancestors_Checked(parent)
                break
        else:
            parent.setCheckState(Qt.Checked)
            self.update_all_ancestors_Checked(parent)

    def update_all_ancestors_Unhecked(self, item):
        parent = item.parent()
        if parent is None:
            return

        tristate = 1
        for child in SelectTab.children_of_(parent):
            if not child.checkState() == Qt.Unchecked:
                parent.setCheckState(tristate)
                self.update_all_ancestors_Unhecked(parent)
                break
        else:
            parent.setCheckState(Qt.Unchecked)
            self.update_all_ancestors_Unhecked(parent)

    @staticmethod
    def children_of_(item):
        for index in range(item.rowCount()):
            child = item.child(index, 0)
            yield child

    def check_all_descendants(self, item):
        for child in SelectTab.children_of_(item):
            if not child.checkState() == Qt.Checked:
                child.setCheckState(Qt.Checked)
                self.check_all_descendants(child)

    def uncheck_all_descendants(self, item):
        for child in SelectTab.children_of_(item):
            if not child.checkState() == Qt.Unchecked:
                child.setCheckState(Qt.Unchecked)
                self.uncheck_all_descendants(child)

    # def initialize_select_button_box(self):
    #     buttonBox = QDialogButtonBox()
    #     buttonBox.setObjectName(u"buttonBox")
    #     buttonBox.setOrientation(Qt.Horizontal)
    #     buttonBox.setStandardButtons(
    #         QDialogButtonBox.Cancel | QDialogButtonBox.Ok)
    #     # buttonBox.accepted.connect(lambda: None)
    #     buttonBox.rejected.connect(sys.exit)
    #     return buttonBox


class Print_docx:
    def __init__(self, list_of_files, filename = "Dokumenty do druku.docx"):
        self.list_of_files = list_of_files
        self.document = Document()
        self.document.add_heading('Dokumenty do przyniesienia', 0)
        self.table = self.make_table()
        self.populate_table()
        self.adjust_column_widths()
        self.document.save(filename)

    def make_table(self):
        table = self.document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nr.'
        hdr_cells[1].text = 'tresc'
        hdr_cells[2].text = 'Znak'
        return table

    def populate_table(self):
        for index, text in enumerate(self.list_of_files, 1):
            row_cells = self.table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = text
            row_cells[2].text = " "

    def adjust_column_widths(self):
        w_nr = 1 / 5
        w_z = 1 / 3.6
        w = self.table.columns[0].width
        self.table.columns[0].width = int(w * w_nr)
        w = self.table.columns[1].width
        self.table.columns[1].width = int(w * 2.5)
        w = self.table.columns[2].width
        self.table.columns[2].width = int(w * w_z)